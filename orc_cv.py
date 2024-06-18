import requests
from docx import Document
from docx.shared import Inches

def get_altmetric_score(doi):
    url = f"https://api.altmetric.com/v1/doi/{doi}"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        return data.get("score", "N/A")
    return "N/A"

def get_citation_count(doi):
    url = f"https://api.crossref.org/works/{doi}"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        return data["message"].get("is-referenced-by-count", 0)
    return 0

def get_oa_status(doi):
    url = f"https://api.unpaywall.org/v2/{doi}?email=your_email@example.com"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        return data.get("is_oa", False)
    return False

def get_authors_from_crossref(doi):
    url = f"https://api.crossref.org/works/{doi}"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        if "author" in data["message"]:
            authors = [author.get("given", "") + " " + author.get("family", "") for author in data["message"]["author"]]
            return ", ".join(authors)
    return ""

def get_orcid_publications(orcid_id):
    url = f"https://pub.orcid.org/v3.0/{orcid_id}/works"
    headers = {"Accept": "application/json"}
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        data = response.json()
        publications = []
        for work in data["group"]:
            work_summary = work["work-summary"][0]
            doi = next((external_id["external-id-value"] for external_id in work["external-ids"]["external-id"] if external_id["external-id-type"] == "doi"), None)
            authors = get_authors_from_crossref(doi) if doi else ""
            publication = {
                "title": work_summary["title"]["title"]["value"],
                "authors": authors,
                "journal": work_summary.get("journal-title", {}).get("value", "") if work_summary.get("journal-title") else "",
                "year": work_summary.get("publication-date", {}).get("year", {}).get("value", ""),
                "doi": doi
            }
            publications.append(publication)
        return publications
    return []

def get_orcid_personal_info(orcid_id):
    url = f"https://pub.orcid.org/v3.0/{orcid_id}/person"
    headers = {"Accept": "application/json"}
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        data = response.json()
        name = data["name"]["given-names"]["value"] + " " + data["name"]["family-name"]["value"]
        email = next((email["email"] for email in data["emails"]["email"] if email["primary"]), "")
        orcid_url = f"https://orcid.org/{orcid_id}"
        return name, email, orcid_url
    return "", "", ""

def get_orcid_education(orcid_id):
    url = f"https://pub.orcid.org/v3.0/{orcid_id}/educations"
    headers = {"Accept": "application/json"}
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        data = response.json()
        educations = []
        for education in data["affiliation-group"]:
            for summary in education["summaries"]:
                education_item = f"{summary['education-summary']['role-title']} in {summary['education-summary']['department-name']}, {summary['education-summary']['organization']['name']}, {summary['education-summary']['end-date']['year']['value']}"
                educations.append(education_item)
        return educations
    return []

# Replace 'your_orcid_id' with the actual ORCID ID
orcid_id = "0000-0002-1360-4052"
email = "aidan@cornelius-bell.com"
publications = get_orcid_publications(orcid_id)
name, orcid_url = get_orcid_personal_info(orcid_id)
educations = get_orcid_education(orcid_id)

# Create a new Word document
document = Document()

# Add a title
document.add_heading("Academic CV", 0)

# Add personal information
document.add_heading("Personal Information", level=1)
document.add_paragraph(f"Name: {name}")
document.add_paragraph(f"Email: {email}")
document.add_paragraph(f"ORCID: {orcid_url}")

# Add education section
document.add_heading("Education", level=1)
for education in educations:
    document.add_paragraph(education)

# Add publications section
document.add_heading("Publications", level=1)

# Create a table with headings
table = document.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Title'
hdr_cells[1].text = 'Authors'
hdr_cells[2].text = 'Journal'
hdr_cells[3].text = 'Year'
hdr_cells[4].text = 'Additional Information'

# Populate the table with publications
for pub in publications:
    row_cells = table.add_row().cells
    row_cells[0].text = pub["title"]
    row_cells[1].text = pub["authors"]
    row_cells[2].text = pub["journal"]
    row_cells[3].text = pub["year"]

    doi = pub["doi"]
    if doi:
        altmetric_score = get_altmetric_score(doi)
        citation_count = get_citation_count(doi)
        is_open_access = get_oa_status(doi)

        additional_info = f"DOI: {doi}\nAltmetric Score: {altmetric_score}\nCitation Count: {citation_count}\nOpen Access: {'Yes' if is_open_access else 'No'}"
        row_cells[4].text = additional_info
        
# Save the document
document.save("academic_cv.docx")

